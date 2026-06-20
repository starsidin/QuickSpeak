import os
import time
from datetime import datetime
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QTextEdit, QCheckBox,
    QSystemTrayIcon, QMenu, QFileDialog, QMessageBox, QSizeGrip,
    QComboBox, QStyle, QProgressBar, QSizePolicy
)
from PySide6.QtCore import Qt, QPoint, QTimer
from PySide6.QtGui import QIcon, QAction

from config import (
    DOUBAO_API_KEY, TEMP_WAV_FILE, LOCAL_BACKEND_URL, ASR_PROVIDER, LOCAL_MODEL,
    SAVE_FOLDER, HOTKEY_ENABLED, HOTKEY_KEY, MICROPHONE_DEVICE, save_settings,
    APP_NAME, APP_DISPLAY_NAME, APP_VERSION, APP_AUTHOR, TEMP_DIR, AUTO_SAVE_FILE
)
from clipboard_util import copy_to_clipboard
from audio_recorder import AudioRecorder
from asr_client import DoubaoASRClient, LocalASRClient
from settings_dialog import SettingsDialog
from PySide6.QtCore import Signal

class StatusLabel(QLabel):
    """状态标签，setText 时自动同步 tooltip"""
    def setText(self, text):
        super().setText(text)
        self.setToolTip(text)


class MainWindow(QMainWindow):
    # Qt signals for hotkey cross-thread communication
    hotkey_pressed = Signal()
    hotkey_released = Signal()

    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"{APP_DISPLAY_NAME} QuickSpeak v{APP_VERSION}")
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.FramelessWindowHint | Qt.Window)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.resize(350, 180)
        self.old_pos = None

        self.record_start_time = 0
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_timer)

        # 自动保存定时器
        self.auto_save_timer = QTimer(self)
        self.auto_save_timer.timeout.connect(self._auto_save_content)
        self.auto_save_timer.start(30000)  # 每30秒自动保存一次

        # 清理临时文件（启动时）
        self._cleanup_temp_files()

        # 加载自动保存的内容
        self._load_auto_saved_content()

        # 获取麦克风设备索引
        mic_device = self._get_microphone_device_index()
        self.recorder = AudioRecorder(TEMP_WAV_FILE, device_index=mic_device)
        
        self._init_asr_client()
        
        self.is_creative_mode = False
        self.last_transcription_audio_path = None

        # 设置相关
        self.save_folder = SAVE_FOLDER  # 用户自定义保存目录，空则用文档目录
        self.hotkey_enabled = HOTKEY_ENABLED
        self.hotkey_key = HOTKEY_KEY
        self._hotkey_listener = None
        self._hotkey_pressed = False  # 防止重复触发

        self.init_ui()
        self.init_tray()
        self.connect_signals()

        # 启动全局快捷键监听
        if self.hotkey_enabled and self.hotkey_key:
            self._start_hotkey_listener()

        self.status_label.setText("检查 API 连接...")
        self.asr.check_health()

    def _get_microphone_device_index(self):
        """获取配置的麦克风设备索引"""
        if not MICROPHONE_DEVICE:
            return None
        try:
            return int(MICROPHONE_DEVICE)
        except (ValueError, TypeError):
            # 如果是设备名称，查找对应索引
            devices = AudioRecorder.get_input_devices()
            for dev in devices:
                if MICROPHONE_DEVICE in dev['name']:
                    return dev['index']
        return None

    def _cleanup_temp_files(self):
        """清理临时文件（启动时、退出时调用）"""
        import glob
        try:
            # 清理临时录音文件
            if os.path.exists(TEMP_WAV_FILE):
                os.remove(TEMP_WAV_FILE)
            # 清理其他可能的临时文件（保留自动保存文件）
            for pattern in ['*.wav', '*.mp3', '*.m4a', '*.flac']:
                for f in glob.glob(os.path.join(TEMP_DIR, pattern)):
                    try:
                        os.remove(f)
                    except:
                        pass
        except Exception as e:
            print(f"清理临时文件失败: {e}")

    def _auto_save_content(self):
        """自动保存当前文本内容"""
        try:
            text = self.text_asr.toPlainText()
            if text.strip():
                with open(AUTO_SAVE_FILE, 'w', encoding='utf-8') as f:
                    f.write(text)
        except Exception as e:
            print(f"自动保存失败: {e}")

    def _load_auto_saved_content(self):
        """加载自动保存的内容"""
        try:
            if os.path.exists(AUTO_SAVE_FILE):
                with open(AUTO_SAVE_FILE, 'r', encoding='utf-8') as f:
                    content = f.read()
                if content.strip():
                    self.text_asr.setText(content)
                    self.status_label.setText("已恢复上次未保存的内容")
                    self.status_label.setStyleSheet("color: #FF9800;")
        except Exception as e:
            print(f"加载自动保存内容失败: {e}")

    def _init_asr_client(self):
        if ASR_PROVIDER == "local":
            self.asr = LocalASRClient(backend_url=LOCAL_BACKEND_URL, model=LOCAL_MODEL)
        else:
            self.asr = DoubaoASRClient(api_key=DOUBAO_API_KEY)

    def _switch_asr_client(self, provider, model=None):
        old_asr = self.asr
        self._disconnect_asr_signals(old_asr)
        
        if provider == "local":
            self.asr = LocalASRClient(backend_url=LOCAL_BACKEND_URL, model=model or LOCAL_MODEL)
        else:
            self.asr = DoubaoASRClient(api_key=DOUBAO_API_KEY)
        
        self._connect_asr_signals()
        self.asr.check_health()

    def _disconnect_asr_signals(self, asr_client):
        try:
            asr_client.backend_status.disconnect(self.handle_backend_status)
        except Exception:
            pass
        try:
            asr_client.request_started.disconnect(self.handle_asr_started)
        except Exception:
            pass
        try:
            asr_client.request_finished.disconnect(self.handle_asr_finished)
        except Exception:
            pass
        try:
            asr_client.request_failed.disconnect(self.handle_error)
        except Exception:
            pass

    def _update_mode_label(self):
        """更新识别模式和云端状态标签"""
        if ASR_PROVIDER == "local":
            mode_text = "本地识别"
            mode_color = "#4CAF50"  # 绿色表示本地
        else:
            mode_text = "云端识别"
            mode_color = "#FF9800"  # 橙色表示云端
        
        self.mode_label.setText(mode_text)
        self.mode_label.setToolTip("不上传云端" if ASR_PROVIDER == "local" else "音频上传至豆包API")
        self.mode_label.setStyleSheet(f"font-size: 11px; color: white; padding: 2px 6px; background-color: {mode_color}; border-radius: 3px;")

    def _set_status(self, text, color="#888888"):
        """更新状态标签的文本和 tooltip"""
        self.status_label.setText(text)
        self.status_label.setToolTip(text)
        self.status_label.setStyleSheet(f"font-size: 12px; color: {color};")

    def init_ui(self):
        self.central_widget = QWidget(self)
        self.central_widget.setStyleSheet("""
            QWidget#CentralWidget {
                background-color: #f5f9ff;
                border-radius: 12px;
            }
            QLabel { color: #555555; font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; }
            QPushButton {
                background-color: #4A90E2;
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: bold;
                font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            }
            QPushButton:hover { background-color: #357ABD; }
            QPushButton:pressed { background-color: #285A8C; }
            QPushButton:disabled { background-color: #cccccc; color: #888888; }
            QTextEdit {
                background-color: #ffffff;
                color: #333333;
                border: 1px solid #dcdcdc;
                border-radius: 8px;
                padding: 8px;
                font-size: 13px;
                font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
                placeholder-text-color: #bbbbbb;
            }
            QCheckBox {
                color: #555555;
                font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            }
        """)
        self.central_widget.setObjectName("CentralWidget")
        self.setCentralWidget(self.central_widget)

        main_layout = QVBoxLayout(self.central_widget)
        main_layout.setContentsMargins(6, 4, 6, 4)
        main_layout.setSpacing(3)

        top_layout = QHBoxLayout()
        title_label = QLabel(APP_DISPLAY_NAME)
        title_label.setStyleSheet("font-size: 15px; font-weight: bold; color: #333333;")
        
        self.status_label = StatusLabel("正在初始化...")
        self.status_label.setFixedHeight(18)
        self.status_label.setMinimumWidth(80)
        self.status_label.setMaximumWidth(120)
        self.status_label.setStyleSheet("font-size: 12px; color: #888888;")
        self.status_label.setWordWrap(False)
        self.status_label.setTextFormat(Qt.PlainText)
        self.status_label.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self.status_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        
        # 识别模式和云端状态标签
        self.mode_label = QLabel("")
        self.mode_label.setStyleSheet("font-size: 11px; color: #666666; padding: 2px 6px; background-color: #f0f0f0; border-radius: 3px;")
        self._update_mode_label()
        
        # 隐藏的 combo_provider，由设置面板控制
        self.combo_provider = QComboBox()
        self.combo_provider.addItem("豆包", "doubao")
        self.combo_provider.addItem("本地千问", "qwen")
        self.combo_provider.hide()
        if ASR_PROVIDER == "local":
            self.combo_provider.setCurrentIndex(1)
        else:
            self.combo_provider.setCurrentIndex(0)

        self.timer_label = QLabel("00:00")
        self.timer_label.setStyleSheet("font-size: 12px; font-weight: bold; color: #f44336;")
        self.timer_label.hide()

        self.btn_settings = QPushButton("设置")
        self.btn_settings.setStyleSheet("""
            QPushButton {
                background-color: #e8f0fe;
                color: #4A90E2;
                font-size: 11px;
                font-weight: bold;
                padding: 2px 10px;
                border: 1px solid #4A90E2;
                border-radius: 4px;
            }
            QPushButton:hover { background-color: #4A90E2; color: white; }
        """)
        self.btn_settings.setFixedHeight(24)
        self.btn_settings.setToolTip("打开设置面板")

        min_btn = QPushButton("—")
        min_btn.setStyleSheet("""
            QPushButton { background-color: transparent; color: #999999; font-size: 16px; padding: 0; font-weight: bold; }
            QPushButton:hover { color: #333333; }
        """)
        min_btn.setFixedSize(20, 20)
        min_btn.clicked.connect(self.showMinimized)

        close_btn = QPushButton("×")
        close_btn.setStyleSheet("""
            QPushButton { background-color: transparent; color: #999999; font-size: 16px; padding: 0; }
            QPushButton:hover { color: #f44336; }
        """)
        close_btn.setFixedSize(20, 20)
        close_btn.clicked.connect(self.hide)

        top_layout.addWidget(title_label)
        top_layout.addWidget(self.mode_label)
        top_layout.addStretch()
        top_layout.addWidget(self.timer_label)
        top_layout.addWidget(self.status_label)
        top_layout.addWidget(self.btn_settings)
        top_layout.addWidget(min_btn)
        top_layout.addWidget(close_btn)
        
        top_container = QWidget()
        top_container.setLayout(top_layout)
        top_container.setContentsMargins(0, 0, 15, 0)
        main_layout.addWidget(top_container)

        btn_container_layout = QVBoxLayout()
        btn_container_layout.setContentsMargins(0, 0, 15, 0)
        btn_container_layout.setSpacing(5)

        btn_layout_row1 = QHBoxLayout()
        self.btn_record_start = QPushButton("开始录音")
        self.btn_record_stop = QPushButton("停止录音")
        self.btn_record_stop.setEnabled(False)
        
        btn_layout_row1.addWidget(self.btn_record_start)
        btn_layout_row1.addWidget(self.btn_record_stop)
        
        btn_layout_row2 = QHBoxLayout()
        self.btn_record_cancel = QPushButton("取消录音")
        self.btn_record_cancel.setEnabled(False)
        self.btn_record_cancel.setStyleSheet("""
            QPushButton { background-color: #ff9800; }
            QPushButton:hover { background-color: #e68a00; }
            QPushButton:disabled { background-color: #cccccc; }
        """)
        
        self.btn_import = QPushButton("导入音频")
        self.btn_save_audio = QPushButton("保存录音")
        self.btn_save_audio.setEnabled(False)
        
        btn_layout_row2.addWidget(self.btn_record_cancel)
        btn_layout_row2.addWidget(self.btn_import)
        btn_layout_row2.addWidget(self.btn_save_audio)
        
        btn_container_layout.addLayout(btn_layout_row1)
        btn_container_layout.addLayout(btn_layout_row2)
        
        btn_container = QWidget()
        btn_container.setLayout(btn_container_layout)
        main_layout.addWidget(btn_container)

        # 音量波形显示
        self.volume_widget = QWidget()
        self.volume_widget.setFixedHeight(40)
        self.volume_widget.setStyleSheet("background-color: #f0f0f0; border-radius: 4px;")
        self.volume_widget.hide()
        
        volume_layout = QVBoxLayout(self.volume_widget)
        volume_layout.setContentsMargins(5, 5, 5, 5)
        volume_layout.setSpacing(0)
        
        self.volume_bar = QProgressBar()
        self.volume_bar.setRange(0, 100)
        self.volume_bar.setValue(0)
        self.volume_bar.setTextVisible(False)
        self.volume_bar.setFixedHeight(30)
        self.volume_bar.setStyleSheet("""
            QProgressBar {
                border: 1px solid #dcdcdc;
                border-radius: 4px;
                background-color: #ffffff;
            }
            QProgressBar::chunk {
                background-color: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                                                  stop:0 #4CAF50, stop:0.5 #FFC107, stop:1 #F44336);
                border-radius: 3px;
            }
        """)
        
        self.volume_label = QLabel("音量: 0%")
        self.volume_label.setStyleSheet("font-size: 10px; color: #666666;")
        self.volume_label.setAlignment(Qt.AlignCenter)
        
        volume_layout.addWidget(self.volume_bar)
        volume_layout.addWidget(self.volume_label)
        
        main_layout.addWidget(self.volume_widget)

        self.text_asr = QTextEdit()
        self.text_asr.setPlaceholderText("识别结果将显示在这里，可直接编辑...")
        self.text_asr.setReadOnly(False)
        self.text_asr.setFixedHeight(70)
        
        result_header_layout = QHBoxLayout()
        lbl_result = QLabel("识别结果:")
        
        self.btn_clear = QPushButton("清空")
        self.btn_clear.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                color: #888888;
                font-weight: normal;
                font-size: 12px;
                padding: 0px 5px;
            }
            QPushButton:hover { color: #f44336; }
        """)

        self.btn_retry = QPushButton("重试识别")
        self.btn_retry.setEnabled(False)
        self.btn_retry.hide()
        self.btn_retry.setStyleSheet("""
            QPushButton {
                background-color: #ff9800;
                color: white;
                font-weight: normal;
                font-size: 12px;
                padding: 2px 8px;
            }
            QPushButton:hover { background-color: #e68a00; }
            QPushButton:disabled { background-color: #cccccc; color: #888888; }
        """)
        
        result_header_layout.addWidget(lbl_result)
        result_header_layout.addStretch()
        result_header_layout.addWidget(self.btn_retry)
        result_header_layout.addWidget(self.btn_clear)
        
        result_header_container = QWidget()
        result_header_container.setLayout(result_header_layout)
        result_header_container.setContentsMargins(0, 0, 15, 0)
        main_layout.addWidget(result_header_container)
        
        self.format_toolbar_widget = QWidget()
        format_layout = QHBoxLayout(self.format_toolbar_widget)
        format_layout.setContentsMargins(0, 0, 0, 5)
        
        format_btn_style = """
            QPushButton {
                font-family: serif; 
                color: #ffffff;
                background-color: #5C9FE6;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #4A90E2;
            }
            QPushButton:checked {
                background-color: #285A8C;
                border: 1px solid #1a3d61;
            }
        """
        
        self.btn_bold = QPushButton("B")
        self.btn_bold.setStyleSheet(format_btn_style + "font-weight: bold;")
        self.btn_bold.setFixedSize(28, 28)
        self.btn_bold.setCheckable(True)
        
        self.btn_italic = QPushButton("I")
        self.btn_italic.setStyleSheet(format_btn_style + "font-style: italic;")
        self.btn_italic.setFixedSize(28, 28)
        self.btn_italic.setCheckable(True)
        
        self.btn_underline = QPushButton("U")
        self.btn_underline.setStyleSheet(format_btn_style + "text-decoration: underline;")
        self.btn_underline.setFixedSize(28, 28)
        self.btn_underline.setCheckable(True)
        
        self.combo_font_size = QComboBox()
        font_sizes = ["8", "9", "10", "11", "12", "13", "14", "16", "18", "20", "24", "28", "32", "36", "48", "72"]
        self.combo_font_size.addItems(font_sizes)
        self.combo_font_size.setCurrentText("13")
        self.combo_font_size.setFixedWidth(60)
        self.combo_font_size.setStyleSheet("""
            QComboBox {
                background-color: #ffffff;
                color: #333333;
                border: 1px solid #dcdcdc;
                border-radius: 4px;
                padding: 2px 5px;
            }
        """)
        
        self.btn_import_text = QPushButton("导入文本")
        self.btn_import_text.setStyleSheet("background-color: #FF9800;")
        
        self.btn_save_txt = QPushButton("保存为 TXT")
        self.btn_save_txt.setStyleSheet("background-color: #4CAF50;")
        
        self.btn_save_word = QPushButton("保存为 Word")
        self.btn_save_word.setStyleSheet("background-color: #2196F3;")
        
        self.btn_exit_creative = QPushButton("退出创作")
        self.btn_exit_creative.setStyleSheet("background-color: #f44336;")
        
        format_layout.addWidget(self.btn_bold)
        format_layout.addWidget(self.btn_italic)
        format_layout.addWidget(self.btn_underline)
        format_layout.addWidget(QLabel(" 字号:"))
        format_layout.addWidget(self.combo_font_size)
        format_layout.addStretch()
        format_layout.addWidget(self.btn_import_text)
        format_layout.addWidget(self.btn_save_txt)
        format_layout.addWidget(self.btn_save_word)
        format_layout.addWidget(self.btn_exit_creative)
        
        self.format_toolbar_widget.hide()

        text_container = QWidget()
        text_layout = QVBoxLayout(text_container)
        text_layout.setContentsMargins(0, 0, 15, 0)
        text_layout.addWidget(self.format_toolbar_widget)
        text_layout.addWidget(self.text_asr)
        main_layout.addWidget(text_container)

        bottom_layout = QHBoxLayout()
        
        self.btn_copy = QPushButton("复制")
        self.btn_copy.setStyleSheet("background-color: #607D8B;")
        self.btn_creative = QPushButton("创作模式")
        self.btn_creative.setStyleSheet("background-color: #9C27B0;")
        
        self.chk_auto_copy = QCheckBox("自动复制")
        self.chk_auto_copy.setChecked(True)
        
        bottom_layout.addWidget(self.btn_copy)
        bottom_layout.addWidget(self.btn_creative)
        bottom_layout.addStretch()
        bottom_layout.addWidget(self.chk_auto_copy)
        
        bottom_container = QWidget()
        bottom_container.setLayout(bottom_layout)
        bottom_container.setContentsMargins(0, 0, 0, 0)
        main_layout.addWidget(bottom_container)
        
        hint_label = QLabel("提示：首次进入或长时间未用后，加载模型需要等待20s。")
        hint_label.setStyleSheet("font-size: 10px; color: #999999;")
        hint_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        hint_layout = QHBoxLayout()
        hint_layout.setContentsMargins(0, 0, 0, 0)
        hint_layout.addStretch()
        hint_layout.addWidget(hint_label)
        
        hint_container = QWidget()
        hint_container.setLayout(hint_layout)
        main_layout.addWidget(hint_container)
        
        self.size_grip = QSizeGrip(self.central_widget)
        self.size_grip.setFixedSize(15, 15)
        self.size_grip.setStyleSheet("background-color: transparent;")
        self.size_grip.move(self.width() - 15, self.height() - 15)
        self.size_grip.show()
        
    def resizeEvent(self, event):
        super().resizeEvent(event)
        if hasattr(self, 'size_grip') and self.size_grip:
            self.size_grip.move(self.width() - 15, self.height() - 15)

    def init_tray(self):
        self.tray_icon = QSystemTrayIcon(self)
        
        # 使用应用程序图标作为系统托盘图标
        icon = QApplication.windowIcon()
        if icon.isNull():
            # 如果没有设置图标，使用默认图标
            icon = QApplication.style().standardIcon(QStyle.SP_ComputerIcon)
        self.tray_icon.setIcon(icon)
        
        tray_menu = QMenu()
        show_action = QAction("显示窗口", self)
        show_action.triggered.connect(self.showNormal)
        quit_action = QAction("退出", self)
        quit_action.triggered.connect(self.quit_app)
        
        tray_menu.addAction(show_action)
        tray_menu.addAction(quit_action)
        
        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.activated.connect(self.tray_icon_activated)
        self.tray_icon.show()

    def tray_icon_activated(self, reason):
        if reason == QSystemTrayIcon.DoubleClick:
            self.showNormal()
            self.activateWindow()

    def connect_signals(self):
        self.btn_record_start.clicked.connect(self.on_record_start)
        self.btn_record_stop.clicked.connect(self.on_record_stop)
        self.btn_record_cancel.clicked.connect(self.on_record_cancel)
        self.btn_import.clicked.connect(self.on_import_audio)
        self.btn_save_audio.clicked.connect(self.on_save_audio)
        self.btn_retry.clicked.connect(self.on_retry_transcription)
        self.btn_clear.clicked.connect(self.on_clear)
        self.btn_copy.clicked.connect(self.on_copy)
        self.btn_settings.clicked.connect(self.open_settings)
        
        self.btn_creative.clicked.connect(self.enter_creative_mode)
        self.btn_exit_creative.clicked.connect(self.exit_creative_mode)
        self.btn_import_text.clicked.connect(self.on_import_text)
        self.btn_save_txt.clicked.connect(self.on_save_txt)
        self.btn_save_word.clicked.connect(self.on_save_word)
        
        self.btn_bold.clicked.connect(self.toggle_bold)
        self.btn_italic.clicked.connect(self.toggle_italic)
        self.btn_underline.clicked.connect(self.toggle_underline)
        self.combo_font_size.currentTextChanged.connect(self.change_font_size)
        self.text_asr.cursorPositionChanged.connect(self.update_format_buttons)
        
        self.recorder.recording_started.connect(self.handle_recording_started)
        self.recorder.recording_stopped.connect(self.handle_recording_stopped)
        self.recorder.recording_canceled.connect(self.handle_recording_canceled)
        self.recorder.error_occurred.connect(self.handle_error)
        self.recorder.volume_level.connect(self.update_volume_display)
        
        self._connect_asr_signals()

    def _connect_asr_signals(self):
        self.asr.backend_status.connect(self.handle_backend_status)
        self.asr.request_started.connect(self.handle_asr_started)
        self.asr.request_finished.connect(self.handle_asr_finished)
        self.asr.request_failed.connect(self.handle_error)


    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            # 检查点击位置是否在某个按钮上，如果是则不拦截
            child = self.childAt(event.position().toPoint())
            if child and isinstance(child, QPushButton):
                event.ignore()
                return
            self.old_pos = event.globalPosition().toPoint()

    def mouseMoveEvent(self, event):
        if self.old_pos is not None:
            delta = event.globalPosition().toPoint() - self.old_pos
            self.move(self.pos() + delta)
            self.old_pos = event.globalPosition().toPoint()

    def mouseReleaseEvent(self, event):
        self.old_pos = None

    def on_record_start(self):
        if os.path.exists(TEMP_WAV_FILE):
            try:
                os.remove(TEMP_WAV_FILE)
            except Exception as e:
                print(f"删除临时文件失败: {e}")
                
        self.btn_save_audio.setEnabled(False)
        self.record_start_time = time.time()
        self.timer_label.setText("00:00")
        self.timer_label.show()
        self.timer.start(1000)
        self.recorder.start_recording()

        self.last_transcription_audio_path = None
        self.btn_retry.hide()
        self.btn_retry.setEnabled(False)

    def on_record_stop(self):
        self.timer.stop()
        self.timer_label.hide()
        self.recorder.stop_recording()
        self.btn_save_audio.setEnabled(True)

    def on_record_cancel(self):
        self.timer.stop()
        self.timer_label.hide()
        self.recorder.cancel_recording()

    def update_timer(self):
        elapsed = int(time.time() - self.record_start_time)
        mins, secs = divmod(elapsed, 60)
        self.timer_label.setText(f"{mins:02d}:{secs:02d}")

    def on_import_audio(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择音频文件",
            "",
            "Audio Files (*.wav *.mp3 *.flac *.ogg *.m4a);;All Files (*)"
        )
        if file_path:
            self.status_label.setText("正在识别导入的音频...")
            self.status_label.setStyleSheet("color: #2196F3;")
            self.last_transcription_audio_path = file_path
            self.btn_retry.hide()
            self.btn_retry.setEnabled(False)
            self.asr.transcribe(file_path)

    def on_save_audio(self):
        if not os.path.exists(TEMP_WAV_FILE):
            self.handle_error("未找到录音文件")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存录音",
            self._get_default_save_path("录音", "wav"),
            "Audio Files (*.wav);;All Files (*)"
        )
        if file_path:
            try:
                import shutil
                shutil.copy2(TEMP_WAV_FILE, file_path)
                self.status_label.setText("录音已保存")
                self.status_label.setStyleSheet("color: #4CAF50;")
            except Exception as e:
                self.handle_error(f"保存录音失败: {str(e)}")

    def on_clear(self):
        self.text_asr.clear()
        self.status_label.setText("已清空")

    def on_retry_transcription(self):
        if not self.last_transcription_audio_path:
            self.handle_error("没有可重试的音频")
            return

        if not self.last_transcription_audio_path.startswith(("http://", "https://")):
            if not os.path.exists(self.last_transcription_audio_path):
                self.handle_error("上一次录音文件已不存在，无法重试")
                return

        self.status_label.setText("正在重试识别...")
        self.status_label.setStyleSheet("color: #2196F3;")
        self.btn_retry.setEnabled(False)
        self.asr.transcribe(self.last_transcription_audio_path)

    def on_copy(self):
        text = self.text_asr.toPlainText()
        if copy_to_clipboard(text):
            self.status_label.setText("已复制")
            self.status_label.setStyleSheet("color: #4CAF50;")

    def enter_creative_mode(self):
        self.is_creative_mode = True
        self.format_toolbar_widget.show()
        self.btn_creative.hide()
        self.btn_copy.hide()
        self.chk_auto_copy.hide()
        self.resize(600, 450)
        
    def exit_creative_mode(self):
        reply = QMessageBox.question(
            self,
            "退出创作",
            "退出创作模式后，当前文本记录将消失，是否确认退出？",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self.is_creative_mode = False
            self.format_toolbar_widget.hide()
            self.btn_creative.show()
            self.btn_copy.show()
            self.chk_auto_copy.show()
            self.text_asr.clear()
            self.resize(350, 250)

    def toggle_bold(self):
        self.text_asr.setFontWeight(700 if self.btn_bold.isChecked() else 400)

    def toggle_italic(self):
        self.text_asr.setFontItalic(self.btn_italic.isChecked())

    def toggle_underline(self):
        self.text_asr.setFontUnderline(self.btn_underline.isChecked())
        
    def change_font_size(self, size_str):
        try:
            size = float(size_str)
            self.text_asr.setFontPointSize(size)
        except ValueError:
            pass

    def update_format_buttons(self):
        self.btn_bold.setChecked(self.text_asr.fontWeight() >= 700)
        self.btn_italic.setChecked(self.text_asr.fontItalic())
        self.btn_underline.setChecked(self.text_asr.fontUnderline())
        
        size = self.text_asr.fontPointSize()
        if size > 0:
            self.combo_font_size.blockSignals(True)
            index = self.combo_font_size.findText(str(int(size)))
            if index >= 0:
                self.combo_font_size.setCurrentIndex(index)
            else:
                self.combo_font_size.setCurrentText(str(int(size)))
            self.combo_font_size.blockSignals(False)

    def on_import_text(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "导入文本文件",
            "",
            "文本文件 (*.txt *.docx);;TXT 文件 (*.txt);;Word 文件 (*.docx);;All Files (*)"
        )
        if not file_path:
            return

        cursor = self.text_asr.textCursor()
        
        try:
            if file_path.lower().endswith(".docx"):
                from docx import Document
                doc = Document(file_path)
                html_content = ""
                for para in doc.paragraphs:
                    para_html = ""
                    for run in para.runs:
                        text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>")
                        
                        pt_size = None
                        if run.font and run.font.size:
                            pt_size = run.font.size.pt
                        elif run.style and run.style.font and run.style.font.size:
                            pt_size = run.style.font.size.pt
                        
                        if pt_size:
                            text = f"<span style='font-size: {pt_size}pt;'>{text}</span>"
                            
                        if run.bold:
                            text = f"<b>{text}</b>"
                        if run.italic:
                            text = f"<i>{text}</i>"
                        if run.underline:
                            text = f"<u>{text}</u>"
                        para_html += text
                        
                    style_name = para.style.name if para.style else ""
                    
                    if style_name.startswith('Heading'):
                        try:
                            level = int(style_name.split(' ')[-1])
                        except:
                            level = 1
                        html_content += f"<h{level}>{para_html}</h{level}>"
                    elif 'List' in style_name:
                        html_content += f"<ul><li>{para_html}</li></ul>"
                    else:
                        if para_html.strip() == "":
                            html_content += "<br>"
                        else:
                            html_content += f"<p>{para_html}</p>"
                
                cursor.insertHtml(html_content)
                
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    text_content = f.read()
                cursor.insertText(text_content)
                
            self.status_label.setText("文本已导入")
            self.status_label.setStyleSheet("color: #4CAF50;")
            
            self.text_asr.setFocus()
            
        except Exception as e:
            self.handle_error(f"导入文本失败: {str(e)}")

    def _get_default_save_path(self, prefix, ext):
        """生成默认保存路径：用户设置目录(或文档目录) + 前缀_时间戳.扩展名"""
        if self.save_folder and os.path.isdir(self.save_folder):
            save_dir = self.save_folder
        else:
            save_dir = os.path.join(os.path.expanduser('~'), 'Documents')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_name = f"{prefix}_{timestamp}.{ext}"
        return os.path.join(save_dir, default_name)

    def on_save_txt(self):
        text = self.text_asr.toPlainText()
        if not text.strip():
            self.handle_error("没有可保存的内容")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存为 TXT 文件",
            self._get_default_save_path("语音识别", "txt"),
            "Text Files (*.txt);;All Files (*)"
        )
        
        if file_path:
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(text)
                self.status_label.setText("已保存为 TXT")
                self.status_label.setStyleSheet("color: #4CAF50;")
            except Exception as e:
                self.handle_error(f"保存 TXT 失败: {str(e)}")

    def on_save_word(self):
        text = self.text_asr.toPlainText()
        if not text.strip():
            self.handle_error("没有可保存的内容")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存为 Word 文件",
            self._get_default_save_path("语音识别", "docx"),
            "Word Files (*.docx);;All Files (*)"
        )
        
        if file_path:
            try:
                from docx import Document
                doc = Document()
                
                text_document = self.text_asr.document()
                for i in range(text_document.blockCount()):
                    block = text_document.findBlockByNumber(i)
                    p = doc.add_paragraph()
                    it = block.begin()
                    while not it.atEnd():
                        fragment = it.fragment()
                        if fragment.isValid():
                            run = p.add_run(fragment.text())
                            char_format = fragment.charFormat()
                            if char_format.fontWeight() >= 700:
                                run.bold = True
                            if char_format.fontItalic():
                                run.italic = True
                            if char_format.fontUnderline():
                                run.underline = True
                        it += 1
                        
                doc.save(file_path)
                self.status_label.setText("已保存为 Word")
                self.status_label.setStyleSheet("color: #4CAF50;")
            except Exception as e:
                self.handle_error(f"保存 Word 失败: {str(e)}")

    def handle_recording_started(self):
        self.status_label.setText("正在录音...")
        self.status_label.setStyleSheet("color: #ff9800;")
        self.btn_record_start.setEnabled(False)
        self.btn_record_stop.setEnabled(True)
        self.btn_record_cancel.setEnabled(True)
        self.btn_exit_creative.setEnabled(False)
        # 显示音量波形
        self.volume_widget.show()
        self.volume_bar.setValue(0)
        self.volume_label.setText("音量: 0%")

    def handle_recording_stopped(self, wav_path):
        self.status_label.setText("录音结束，准备识别...")
        self.btn_record_start.setEnabled(True)
        self.btn_record_stop.setEnabled(False)
        self.btn_record_cancel.setEnabled(False)
        self.btn_exit_creative.setEnabled(True)
        self.last_transcription_audio_path = wav_path
        self.btn_retry.hide()
        self.btn_retry.setEnabled(False)
        # 隐藏音量波形
        self.volume_widget.hide()
        self.volume_bar.setValue(0)
        self.volume_label.setText("音量: 0%")
        self.asr.transcribe(wav_path)
        
    def handle_recording_canceled(self):
        self.status_label.setText("录音已取消")
        self.status_label.setStyleSheet("color: #999999;")
        self.btn_record_start.setEnabled(True)
        self.btn_record_stop.setEnabled(False)
        self.btn_record_cancel.setEnabled(False)
        self.btn_exit_creative.setEnabled(True)
        # 隐藏音量波形
        self.volume_widget.hide()
        self.volume_bar.setValue(0)
        self.volume_label.setText("音量: 0%")

    def update_volume_display(self, level):
        """更新音量显示 (level: 0.0 - 1.0)"""
        percentage = int(level * 100)
        self.volume_bar.setValue(percentage)
        self.volume_label.setText(f"音量: {percentage}%")

    def on_provider_changed(self, index):
        provider_data = self.combo_provider.currentData()
        
        if provider_data == "qwen":
            self._switch_asr_client("local", model="qwen")
            self.status_label.setText("已切换至本地千问")
        else:
            self._switch_asr_client("doubao")
            self.status_label.setText("已切换至豆包")
        
        self.status_label.setStyleSheet("color: #4CAF50;")

    def handle_backend_status(self, is_ok, error_type=""):
        provider_data = self.combo_provider.currentData() if hasattr(self, 'combo_provider') else "doubao"
        
        if is_ok:
            if provider_data == "qwen":
                self.status_label.setText("本地后端连接正常")
            else:
                self.status_label.setText("API 状态正常")
            self.status_label.setStyleSheet("color: #4CAF50;")
        else:
            if error_type == "backend":
                self.status_label.setText("本地后端未启动，请运行 run_backend.bat")
            elif error_type == "api_key":
                self.status_label.setText("API Key 无效，请检查设置")
            elif error_type == "timeout":
                self.status_label.setText("连接超时，请检查网络")
            else:
                if provider_data == "qwen":
                    self.status_label.setText("本地后端未启动，请运行 run_backend.bat")
                else:
                    self.status_label.setText("API 状态异常，请检查 API Key")
            self.status_label.setStyleSheet("color: #f44336;")

    def handle_asr_started(self):
        self.status_label.setText("正在识别...")
        self.status_label.setStyleSheet("color: #2196F3;")

        self.btn_retry.hide()
        self.btn_retry.setEnabled(False)

    def handle_asr_finished(self, text):
        self.status_label.setText("识别完成")
        self.status_label.setStyleSheet("color: #4CAF50;")
        
        self.btn_retry.hide()
        self.btn_retry.setEnabled(False)

        if self.is_creative_mode:
            cursor = self.text_asr.textCursor()
            current_text = self.text_asr.toPlainText()
            if current_text and cursor.position() > 0:
                cursor.insertText("\n" + text)
            else:
                cursor.insertText(text)
        else:
            self.text_asr.setText(text)
        
        self.text_asr.ensureCursorVisible()
        
        if self.chk_auto_copy.isChecked() and not self.is_creative_mode:
            self.on_copy()

    def handle_error(self, err_msg, error_type=""):
        self.status_label.setText("识别失败")
        self.status_label.setToolTip(err_msg)  # 鼠标悬停显示完整错误信息
        self.status_label.setStyleSheet("color: #f44336;")
        can_retry = bool(self.last_transcription_audio_path)
        self.btn_retry.setVisible(can_retry)
        self.btn_retry.setEnabled(can_retry)
        if not self.is_creative_mode:
            retry_hint = "\n\n录音已保留，可以点击'重试识别'。" if can_retry else ""
            QMessageBox.warning(self, "识别失败", f"{err_msg}{retry_hint}")
        print(f"ERROR [{error_type}]: {err_msg}")
        self.btn_record_start.setEnabled(True)
        self.btn_record_stop.setEnabled(False)
        self.btn_record_cancel.setEnabled(False)
        self.btn_exit_creative.setEnabled(True)

    def quit_app(self):
        reply = QMessageBox.question(
            self,
            "退出程序",
            "退出前请确保你需要保存的内容已进行保存。\n确认要退出吗？",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            self._stop_hotkey_listener()
            self.tray_icon.hide()
            # 退出时清理临时文件
            self._cleanup_temp_files()
            QApplication.quit()

    def open_settings(self):
        """打开设置对话框"""
        try:
            current_settings = {
                "save_folder": self.save_folder,
                "hotkey_enabled": self.hotkey_enabled,
                "hotkey_key": self.hotkey_key,
                "asr_provider": ASR_PROVIDER,
                "local_model": LOCAL_MODEL,
            }
            provider_data = self.combo_provider.currentData()
            if provider_data == "doubao":
                current_settings["asr_provider"] = "doubao"
                current_settings["local_model"] = "qwen"
            elif provider_data == "qwen":
                current_settings["asr_provider"] = "local"
                current_settings["local_model"] = "qwen"

            dlg = SettingsDialog(self, current_settings)
            if dlg.exec() == SettingsDialog.Accepted:
                new_settings = dlg.get_settings()
                self._apply_settings(new_settings)
        except Exception as e:
            import traceback
            traceback.print_exc()

    def _apply_settings(self, settings):
        """应用并持久化设置"""
        # 保存目录
        self.save_folder = settings.get("save_folder", "")

        # 快捷键
        old_hotkey_enabled = self.hotkey_enabled
        old_hotkey_key = self.hotkey_key
        self.hotkey_enabled = settings.get("hotkey_enabled", False)
        self.hotkey_key = settings.get("hotkey_key", "")

        if old_hotkey_enabled or old_hotkey_key:
            self._stop_hotkey_listener()
        if self.hotkey_enabled and self.hotkey_key:
            self._start_hotkey_listener()

        # ASR 引擎切换 - 同步到顶部下拉框
        asr_provider = settings.get("asr_provider", "doubao")
        local_model = settings.get("local_model", "qwen")

        self.combo_provider.blockSignals(True)
        if asr_provider == "doubao":
            self.combo_provider.setCurrentIndex(0)
        else:
            self.combo_provider.setCurrentIndex(1)
        self.combo_provider.blockSignals(False)

        # 切换 ASR 客户端
        if asr_provider == "doubao":
            self._switch_asr_client("doubao")
            self.status_label.setText("已切换至豆包")
        else:
            self._switch_asr_client("local", model="qwen")
            self.status_label.setText("已切换至本地千问")
        self.status_label.setStyleSheet("color: #4CAF50;")

        # 持久化
        try:
            save_settings(settings)
        except Exception as e:
            print(f"保存设置失败: {e}")

        self.status_label.setText("设置已保存")
        self.status_label.setStyleSheet("color: #4CAF50;")

    # ---- 全局快捷键（按住说话） ----

    def _key_str_to_pynput(self, key_str):
        """将存储的键名字符串转换为 pynput 可比较的对象"""
        from pynput.keyboard import Key, KeyCode

        special_reverse = {
            "Ctrl(左)": Key.ctrl_l, "Ctrl(右)": Key.ctrl_r,
            "Alt(左)": Key.alt_l, "Alt(右)": Key.alt_r,
            "Shift(左)": Key.shift_l, "Shift(右)": Key.shift_r,
            "空格": Key.space, "Tab": Key.tab,
            "Enter": Key.enter, "退格": Key.backspace,
            "CapsLock": Key.caps_lock, "Esc": Key.esc,
            "F1": Key.f1, "F2": Key.f2, "F3": Key.f3, "F4": Key.f4,
            "F5": Key.f5, "F6": Key.f6, "F7": Key.f7, "F8": Key.f8,
            "F9": Key.f9, "F10": Key.f10, "F11": Key.f11, "F12": Key.f12,
            "Insert": Key.insert, "Delete": Key.delete,
            "Home": Key.home, "End": Key.end,
            "PageUp": Key.page_up, "PageDown": Key.page_down,
            "上": Key.up, "下": Key.down, "左": Key.left, "右": Key.right,
        }

        if key_str in special_reverse:
            return special_reverse[key_str]

        # 普通字母键
        if len(key_str) == 1 and key_str.isalpha():
            return KeyCode.from_char(key_str.lower())

        return None

    def _start_hotkey_listener(self):
        """启动全局键盘监听：按住指定键录音"""
        if not self.hotkey_key:
            return

        target_key = self._key_str_to_pynput(self.hotkey_key)
        if target_key is None:
            print(f"无法识别的快捷键: {self.hotkey_key}")
            return

        from pynput import keyboard

        def on_press(key):
            if key == target_key and not self._hotkey_pressed:
                self._hotkey_pressed = True
                # 使用 Qt Signal 跨线程调用
                self.hotkey_pressed.emit()

        def on_release(key):
            if key == target_key and self._hotkey_pressed:
                self._hotkey_pressed = False
                # 使用 Qt Signal 跨线程调用
                self.hotkey_released.emit()

        self._hotkey_listener = keyboard.Listener(on_press=on_press, on_release=on_release)
        self._hotkey_listener.start()
        
        # 连接信号到槽函数
        self.hotkey_pressed.connect(self._hotkey_start_record)
        self.hotkey_released.connect(self._hotkey_stop_record)

    def _stop_hotkey_listener(self):
        """停止全局键盘监听"""
        if self._hotkey_listener:
            self._hotkey_listener.stop()
            self._hotkey_listener = None
        self._hotkey_pressed = False

    def _hotkey_start_record(self):
        """快捷键触发：开始录音"""
        if self.btn_record_start.isEnabled():
            self.on_record_start()

    def _hotkey_stop_record(self):
        """快捷键触发：停止录音"""
        if self.btn_record_stop.isEnabled():
            self.on_record_stop()
